"""
Streamlit Dashboard & Reporting Module for the Forensic AI Log Analyzer.
"""
from __future__ import annotations

# Standard Library
import io
import json
import tempfile
from dataclasses import asdict, is_dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
import os # Added for os.getenv

# Third-Party
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from loguru import logger

try:
    import networkx as nx
    from pyvis.network import Network
except Exception:  # pragma: no cover
    nx = None
    Network = None

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:  # pragma: no cover
    logger.warning("python-docx not available, Word report generation disabled")
    Document = None

# Local Project Imports (these should exist in your repo)
from data_models import ForensicReport, LogEntry, SeverityLevel  # type: ignore
from agents import ForensicOrchestrator  # type: ignore
from config import Config  # type: ignore
from llm_analysis import GraniteAnalyzer  # type: ignore
from aws_s3_utils import get_s3_utils, is_s3_available  # type: ignore


# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------

def _safe_to_dict(obj: Any) -> Any:
    """Best-effort conversion of complex objects to plain dicts for JSON/reporting."""
    try:
        if is_dataclass(obj):
            return asdict(obj)
        if hasattr(obj, "dict") and callable(getattr(obj, "dict")):
            return obj.dict()  # pydantic-style
        if hasattr(obj, "model_dump"):
            return obj.model_dump()
        if isinstance(obj, (list, tuple)):
            return [_safe_to_dict(x) for x in obj]
        if isinstance(obj, dict):
            return {k: _safe_to_dict(v) for k, v in obj.items()}
        return obj
    except Exception:
        return str(obj)


def _pretty_json(data: Any) -> str:
    try:
        return json.dumps(_safe_to_dict(data), indent=2, default=str)
    except Exception:
        return str(data)


# -----------------------------------------------------------------------------
# Report Generator
# -----------------------------------------------------------------------------
class ReportGenerator:
    """Generates downloadable reports (DOCX/JSON) from a ForensicReport."""

    def generate_docx(self, report: ForensicReport) -> bytes:
        if Document is None:
            raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")

        doc = Document()
        doc.add_heading("Forensic AI Log Analyzer Report", 0)

        # Case Metadata
        p = doc.add_paragraph()
        p.add_run("Case ID: ").bold = True
        p.add_run(str(getattr(report, "case_id", "N/A")))
        p = doc.add_paragraph()
        p.add_run("Generated: ").bold = True
        p.add_run(datetime.utcnow().strftime("%Y-%m-%d %H:%M:%SZ"))

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(str(getattr(report, "executive_summary", "No summary available.")))

        # Key Metrics
        doc.add_heading("Key Metrics", level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Value"
        metrics = {
            "Case ID": getattr(report, "case_id", "N/A"),
            "Total Events": f"{getattr(report, 'total_events', 0):,}",
            "Severity": str(getattr(getattr(report.analysis.initial_analysis, 'severity', ''), 'value', 'N/A')).upper()
            if getattr(report, 'analysis', None) and getattr(report.analysis, 'initial_analysis', None) else "N/A",
            "Confidence": f"{getattr(getattr(report.analysis.initial_analysis, 'confidence', 0.0), '__float__', lambda: report.analysis.initial_analysis.confidence)():.1%}"
            if getattr(report, 'analysis', None) and getattr(report.analysis, 'initial_analysis', None) else "N/A",
        }
        for k, v in metrics.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(k)
            row_cells[1].text = str(v)

        # IOCs
        doc.add_heading("Indicators of Compromise", level=1)
        iocs = getattr(getattr(report.analysis, 'initial_analysis', {}), 'iocs', []) if getattr(report, 'analysis', None) else []
        if iocs:
            t = doc.add_table(rows=1, cols=4)
            t.style = "Light Grid Accent 1"
            t.rows[0].cells[0].text = "Type"
            t.rows[0].cells[1].text = "Value"
            t.rows[0].cells[2].text = "Confidence"
            t.rows[0].cells[3].text = "Context"
            for i in iocs:
                row = t.add_row().cells
                row[0].text = str(getattr(i, 'type', '')).upper()
                row[1].text = str(getattr(i, 'value', ''))
                conf = getattr(i, 'confidence', None)
                row[2].text = f"{float(conf):.1%}" if conf is not None else "N/A"
                row[3].text = str(getattr(i, 'context', ''))
        else:
            doc.add_paragraph("No IOCs identified.")

        # MITRE
        doc.add_heading("MITRE ATT&CK Mapping", level=1)
        mappings = getattr(getattr(report.analysis, 'initial_analysis', {}), 'mitre_mappings', []) if getattr(report, 'analysis', None) else []
        if mappings:
            t = doc.add_table(rows=1, cols=5)
            t.style = "Light Grid"
            t.rows[0].cells[0].text = "Tactic ID"
            t.rows[0].cells[1].text = "Tactic"
            t.rows[0].cells[2].text = "Technique ID"
            t.rows[0].cells[3].text = "Technique"
            t.rows[0].cells[4].text = "Confidence"
            for m in mappings:
                row = t.add_row().cells
                row[0].text = str(getattr(m, 'tactic_id', ''))
                row[1].text = str(getattr(m, 'tactic_name', ''))
                row[2].text = str(getattr(m, 'technique_id', 'N/A'))
                row[3].text = str(getattr(m, 'technique_name', 'N/A'))
                conf = getattr(m, 'confidence', None)
                row[4].text = f"{float(conf):.1%}" if conf is not None else "N/A"
        else:
            doc.add_paragraph("No MITRE ATT&CK mappings identified.")

        # Timeline (compact)
        doc.add_heading("Event Timeline (compact)", level=1)
        timeline = getattr(report, 'timeline', []) or []
        if timeline:
            t = doc.add_table(rows=1, cols=5)
            t.style = "Light Shading"
            t.rows[0].cells[0].text = "Timestamp"
            t.rows[0].cells[1].text = "Source"
            t.rows[0].cells[2].text = "Severity"
            t.rows[0].cells[3].text = "Event Type"
            t.rows[0].cells[4].text = "Message"
            for ev in timeline[:200]:  # cap rows to keep docx reasonable
                row = t.add_row().cells
                row[0].text = str(ev.get('timestamp', ''))
                row[1].text = str(ev.get('source', ''))
                row[2].text = str(ev.get('severity', ''))
                row[3].text = str(ev.get('event_type', ''))
                row[4].text = str(ev.get('message', ''))
        else:
            doc.add_paragraph("No timeline data available.")

        # Recommendations
        doc.add_heading("Recommendations", level=1)
        recs = getattr(getattr(report, 'analysis', {}), 'recommendations', [])
        if recs:
            for r in recs:
                doc.add_paragraph(str(r), style="List Bullet")
        else:
            doc.add_paragraph("No specific recommendations generated.")

        # Serialize
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def generate_json(self, report: ForensicReport) -> bytes:
        return _pretty_json(report).encode("utf-8")

    def generate_html_report(self, report: ForensicReport) -> str:
        """Generate a comprehensive HTML report with embedded charts."""
        html_content = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Forensic Analysis Report - {getattr(report, 'case_id', 'Case')}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 10px; margin-bottom: 30px; }}
                .metric {{ background: #f8f9fa; padding: 20px; border-radius: 8px; margin: 10px 0; border-left: 4px solid #007bff; }}
                .section {{ margin: 30px 0; }}
                .section h2 {{ color: #333; border-bottom: 2px solid #007bff; padding-bottom: 10px; }}
                .ioc-item {{ background: #fff3cd; padding: 15px; margin: 10px 0; border-radius: 5px; border-left: 4px solid #ffc107; }}
                .mitre-item {{ background: #d1ecf1; padding: 15px; margin: 10px 0; border-radius: 5px; border-left: 4px solid #17a2b8; }}
                .recommendation {{ background: #d4edda; padding: 15px; margin: 10px 0; border-radius: 5px; border-left: 4px solid #28a745; }}
                .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; margin: 20px 0; }}
                .chart-container {{ background: white; padding: 20px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); margin: 20px 0; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th, td {{ padding: 12px; text-align: left; border-bottom: 1px solid #ddd; }}
                th {{ background-color: #f2f2f2; font-weight: bold; }}
                .severity-critical {{ color: #dc3545; font-weight: bold; }}
                .severity-high {{ color: #fd7e14; font-weight: bold; }}
                .severity-medium {{ color: #ffc107; font-weight: bold; }}
                .severity-low {{ color: #28a745; font-weight: bold; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>🔍 Forensic AI Log Analyzer Report</h1>
                <p><strong>Case ID:</strong> {getattr(report, 'case_id', 'N/A')}</p>
                <p><strong>Generated:</strong> {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%SZ")}</p>
            </div>

            <div class="section">
                <h2>📋 Executive Summary</h2>
                <div class="metric">
                    {getattr(report, 'executive_summary', 'No summary available.')}
                </div>
            </div>

            <div class="section">
                <h2>📊 Key Metrics</h2>
                <div class="grid">
                    <div class="metric">
                        <h3>Total Events</h3>
                        <p style="font-size: 24px; font-weight: bold; color: #007bff;">{getattr(report, 'total_events', 0):,}</p>
                    </div>
                    <div class="metric">
                        <h3>Severity</h3>
                        <p style="font-size: 24px; font-weight: bold; color: #dc3545;">{str(getattr(getattr(report.analysis.initial_analysis, 'severity', ''), 'value', 'N/A')).upper() if getattr(report, 'analysis', None) and getattr(report.analysis, 'initial_analysis', None) else 'N/A'}</p>
                    </div>
                    <div class="metric">
                        <h3>Confidence</h3>
                        <p style="font-size: 24px; font-weight: bold; color: #28a745;">{f"{getattr(getattr(report.analysis.initial_analysis, 'confidence', 0.0), '__float__', lambda: report.analysis.initial_analysis.confidence)():.1%}" if getattr(report, 'analysis', None) and getattr(report.analysis, 'initial_analysis', None) else 'N/A'}</p>
                    </div>
                </div>
            </div>
        """

        # Add IOCs section
        iocs = getattr(getattr(report.analysis, 'initial_analysis', {}), 'iocs', []) if getattr(report, 'analysis', None) else []
        if iocs:
            html_content += """
            <div class="section">
                <h2>🚨 Indicators of Compromise</h2>
            """
            for ioc in iocs:
                conf = getattr(ioc, 'confidence', None)
                confidence_str = f"{float(conf):.1%}" if conf is not None else "N/A"
                html_content += f"""
                <div class="ioc-item">
                    <h4>{str(getattr(ioc, 'type', '')).upper()}</h4>
                    <p><strong>Value:</strong> {str(getattr(ioc, 'value', ''))}</p>
                    <p><strong>Confidence:</strong> {confidence_str}</p>
                    <p><strong>Context:</strong> {str(getattr(ioc, 'context', ''))}</p>
                </div>
                """
            html_content += "</div>"

        # Add MITRE section
        mappings = getattr(getattr(report.analysis, 'initial_analysis', {}), 'mitre_mappings', []) if getattr(report, 'analysis', None) else []
        if mappings:
            html_content += """
            <div class="section">
                <h2>🎯 MITRE ATT&CK Mapping</h2>
            """
            for m in mappings:
                conf = getattr(m, 'confidence', None)
                confidence_str = f"{float(conf):.1%}" if conf is not None else "N/A"
                html_content += f"""
                <div class="mitre-item">
                    <h4>{getattr(m, 'tactic_name', '')} ({getattr(m, 'tactic_id', '')})</h4>
                    <p><strong>Technique:</strong> {getattr(m, 'technique_name', 'N/A')} ({getattr(m, 'technique_id', 'N/A')})</p>
                    <p><strong>Confidence:</strong> {confidence_str}</p>
                </div>
                """
            html_content += "</div>"

        # Add recommendations
        recs = getattr(getattr(report, 'analysis', {}), 'recommendations', [])
        if recs:
            html_content += """
            <div class="section">
                <h2>💡 Recommendations</h2>
            """
            for i, r in enumerate(recs, 1):
                html_content += f"""
                <div class="recommendation">
                    <p><strong>{i}.</strong> {r}</p>
                </div>
                """
            html_content += "</div>"

        # Add timeline section
        timeline = getattr(report, 'timeline', []) or []
        if timeline:
            html_content += """
            <div class="section">
                <h2>⏰ Event Timeline (Recent Events)</h2>
                <table>
                    <thead>
                        <tr>
                            <th>Timestamp</th>
                            <th>Source</th>
                            <th>Severity</th>
                            <th>Event Type</th>
                            <th>Message</th>
                        </tr>
                    </thead>
                    <tbody>
            """
            for ev in timeline[:50]:  # Show first 50 events
                severity_class = f"severity-{ev.get('severity', 'info')}"
                html_content += f"""
                    <tr>
                        <td>{ev.get('timestamp', '')}</td>
                        <td>{ev.get('source', '')}</td>
                        <td class="{severity_class}">{ev.get('severity', '').upper()}</td>
                        <td>{ev.get('event_type', '')}</td>
                        <td>{ev.get('message', '')[:100]}{'...' if len(str(ev.get('message', ''))) > 100 else ''}</td>
                    </tr>
                """
            html_content += """
                    </tbody>
                </table>
            </div>
            """

        html_content += """
        </body>
        </html>
        """
        
        return html_content


# -----------------------------------------------------------------------------
# Chart Generator
# -----------------------------------------------------------------------------
class ChartGenerator:
    """Generates comprehensive charts and visualizations for forensic analysis."""

    def __init__(self):
        self.color_scheme = {
            'critical': '#FF0000',
            'high': '#FF6B6B', 
            'medium': '#FFA500',
            'low': '#FFD700',
            'info': '#87CEEB',
            'success': '#32CD32',
            'warning': '#FFA500',
            'error': '#DC143C'
        }

    def generate_severity_distribution(self, timeline_data: List[Dict[str, Any]]) -> go.Figure:
        """Generate severity distribution chart."""
        if not timeline_data:
            return self._create_empty_chart("No timeline data available")
        
        # Count severity levels
        severity_counts = {}
        for event in timeline_data:
            severity = event.get('severity', 'unknown').lower()
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        if not severity_counts:
            return self._create_empty_chart("No severity data found")
        
        # Create pie chart
        fig = go.Figure(data=[go.Pie(
            labels=list(severity_counts.keys()),
            values=list(severity_counts.values()),
            hole=0.4,
            marker_colors=[self.color_scheme.get(sev, '#808080') for sev in severity_counts.keys()]
        )])
        
        fig.update_layout(
            title="Event Severity Distribution",
            title_x=0.5,
            showlegend=True,
            height=400,
            margin=dict(t=50, b=50, l=50, r=50)
        )
        
        return fig

    def generate_timeline_heatmap(self, timeline_data: List[Dict[str, Any]]) -> go.Figure:
        """Generate timeline heatmap showing event density over time."""
        if not timeline_data:
            return self._create_empty_chart("No timeline data available")
        
        # Extract timestamps and create hourly bins
        timestamps = []
        for event in timeline_data:
            try:
                ts = pd.to_datetime(event.get('timestamp', ''))
                timestamps.append(ts)
            except:
                continue
        
        if not timestamps:
            return self._create_empty_chart("No valid timestamps found")
        
        # Create hourly bins
        df = pd.DataFrame({'timestamp': timestamps})
        df['hour'] = df['timestamp'].dt.hour
        df['day'] = df['timestamp'].dt.date
        
        # Count events per hour per day
        heatmap_data = df.groupby(['day', 'hour']).size().unstack(fill_value=0)
        
        fig = go.Figure(data=go.Heatmap(
            z=heatmap_data.values,
            x=heatmap_data.columns,
            y=heatmap_data.index,
            colorscale='Reds',
            showscale=True
        ))
        
        fig.update_layout(
            title="Event Timeline Heatmap (Events per Hour)",
            xaxis_title="Hour of Day",
            yaxis_title="Date",
            height=400,
            margin=dict(t=50, b=50, l=50, r=50)
        )
        
        return fig

    def generate_source_analysis(self, timeline_data: List[Dict[str, Any]]) -> go.Figure:
        """Generate source analysis chart showing events by source."""
        if not timeline_data:
            return self._create_empty_chart("No timeline data available")
        
        # Count events by source
        source_counts = {}
        for event in timeline_data:
            source = event.get('source', 'unknown')
            source_counts[source] = source_counts.get(source, 0) + 1
        
        if not source_counts:
            return self._create_empty_chart("No source data found")
        
        # Create bar chart
        sources = list(source_counts.keys())
        counts = list(source_counts.values())
        
        fig = go.Figure(data=[go.Bar(
            x=sources,
            y=counts,
            marker_color='lightblue',
            text=counts,
            textposition='auto'
        )])
        
        fig.update_layout(
            title="Events by Source",
            xaxis_title="Source",
            yaxis_title="Event Count",
            height=400,
            margin=dict(t=50, b=100, l=50, r=50),
            xaxis_tickangle=-45
        )
        
        return fig

    def generate_ioc_analysis(self, iocs: List[Any]) -> go.Figure:
        """Generate IOC analysis chart."""
        if not iocs:
            return self._create_empty_chart("No IOCs found")
        
        # Extract IOC data
        ioc_types = {}
        ioc_confidence = []
        
        for ioc in iocs:
            ioc_type = str(getattr(ioc, 'type', 'unknown')).upper()
            confidence = float(getattr(ioc, 'confidence', 0.0))
            
            ioc_types[ioc_type] = ioc_types.get(ioc_type, 0) + 1
            ioc_confidence.append(confidence)
        
        if not ioc_types:
            return self._create_empty_chart("No valid IOC data found")
        
        # Create subplot with pie chart and histogram
        fig = make_subplots(
            rows=1, cols=2,
            subplot_titles=('IOCs by Type', 'IOC Confidence Distribution'),
            specs=[[{"type": "pie"}, {"type": "histogram"}]]
        )
        
        # Pie chart
        fig.add_trace(
            go.Pie(
                labels=list(ioc_types.keys()),
                values=list(ioc_types.values()),
                hole=0.3
            ),
            row=1, col=1
        )
        
        # Histogram
        fig.add_trace(
            go.Histogram(
                x=ioc_confidence,
                nbinsx=10,
                marker_color='lightgreen'
            ),
            row=1, col=2
        )
        
        fig.update_layout(
            title="IOC Analysis",
            height=400,
            margin=dict(t=50, b=50, l=50, r=50)
        )
        
        return fig

    def generate_mitre_analysis(self, mitre_mappings: List[Any]) -> go.Figure:
        """Generate MITRE ATT&CK analysis chart."""
        if not mitre_mappings:
            return self._create_empty_chart("No MITRE mappings found")
        
        # Extract MITRE data
        tactic_counts = {}
        technique_counts = {}
        confidence_scores = []
        
        for mapping in mitre_mappings:
            tactic = str(getattr(mapping, 'tactic_name', 'Unknown'))
            technique = str(getattr(mapping, 'technique_name', 'Unknown'))
            confidence = float(getattr(mapping, 'confidence', 0.0))
            
            tactic_counts[tactic] = tactic_counts.get(tactic, 0) + 1
            technique_counts[technique] = technique_counts.get(technique, 0) + 1
            confidence_scores.append(confidence)
        
        if not tactic_counts:
            return self._create_empty_chart("No valid MITRE data found")
        
        # Create subplot
        fig = make_subplots(
            rows=2, cols=1,
            subplot_titles=('MITRE Tactics Distribution', 'MITRE Confidence Distribution'),
            vertical_spacing=0.1
        )
        
        # Tactics bar chart
        fig.add_trace(
            go.Bar(
                x=list(tactic_counts.keys()),
                y=list(tactic_counts.values()),
                marker_color='lightcoral',
                text=list(tactic_counts.values()),
                textposition='auto'
            ),
            row=1, col=1
        )
        
        # Confidence histogram
        fig.add_trace(
            go.Histogram(
                x=confidence_scores,
                nbinsx=10,
                marker_color='lightblue'
            ),
            row=2, col=1
        )
        
        fig.update_layout(
            title="MITRE ATT&CK Analysis",
            height=600,
            margin=dict(t=50, b=50, l=50, r=50),
            xaxis_tickangle=-45
        )
        
        return fig

    def generate_event_timeline(self, timeline_data: List[Dict[str, Any]]) -> go.Figure:
        """Generate interactive event timeline."""
        if not timeline_data:
            return self._create_empty_chart("No timeline data available")
        
        # Prepare data
        timestamps = []
        sources = []
        severities = []
        messages = []
        
        for event in timeline_data:
            try:
                ts = pd.to_datetime(event.get('timestamp', ''))
                timestamps.append(ts)
                sources.append(event.get('source', 'unknown'))
                severities.append(event.get('severity', 'info'))
                messages.append(event.get('message', '')[:100] + '...' if len(str(event.get('message', ''))) > 100 else event.get('message', ''))
            except:
                continue
        
        if not timestamps:
            return self._create_empty_chart("No valid timestamps found")
        
        # Create scatter plot
        fig = go.Figure()
        
        # Add traces for each severity level
        severity_levels = ['critical', 'high', 'medium', 'low', 'info']
        for severity in severity_levels:
            mask = [s == severity for s in severities]
            if any(mask):
                fig.add_trace(go.Scatter(
                    x=[timestamps[i] for i in range(len(timestamps)) if mask[i]],
                    y=[sources[i] for i in range(len(sources)) if mask[i]],
                    mode='markers',
                    name=severity.upper(),
                    marker=dict(
                        size=10,
                        color=self.color_scheme.get(severity, '#808080'),
                        opacity=0.7
                    ),
                    text=[messages[i] for i in range(len(messages)) if mask[i]],
                    hovertemplate='<b>%{text}</b><br>Time: %{x}<br>Source: %{y}<br>Severity: ' + severity.upper() + '<extra></extra>'
                ))
        
        fig.update_layout(
            title="Event Timeline",
            xaxis_title="Time",
            yaxis_title="Source",
            height=500,
            margin=dict(t=50, b=50, l=50, r=50),
            hovermode='closest'
        )
        
        return fig

    def generate_attack_chain_visualization(self, attack_chain: List[str]) -> go.Figure:
        """Generate attack chain visualization."""
        if not attack_chain:
            return self._create_empty_chart("No attack chain data available")
        
        # Create sankey diagram for attack flow
        fig = go.Figure(data=[go.Sankey(
            node=dict(
                pad=15,
                thickness=20,
                line=dict(color="black", width=0.5),
                label=attack_chain,
                color="blue"
            ),
            link=dict(
                source=[i for i in range(len(attack_chain)-1)],
                target=[i+1 for i in range(len(attack_chain)-1)],
                value=[1] * (len(attack_chain)-1)
            )
        )])
        
        fig.update_layout(
            title="Attack Chain Flow",
            font_size=10,
            height=400,
            margin=dict(t=50, b=50, l=50, r=50)
        )
        
        return fig

    def generate_comprehensive_dashboard(self, report: ForensicReport) -> List[go.Figure]:
        """Generate comprehensive dashboard with all charts."""
        charts = []
        
        # Get data from report
        timeline_data = getattr(report, 'timeline', []) or []
        iocs = getattr(getattr(report.analysis, 'initial_analysis', {}), 'iocs', []) if getattr(report, 'analysis', None) else []
        mitre_mappings = getattr(getattr(report.analysis, 'initial_analysis', {}), 'mitre_mappings', []) if getattr(report, 'analysis', None) else []
        attack_chain = getattr(getattr(report, 'analysis', {}), 'attack_chain', []) if getattr(report, 'analysis', None) else []
        
        # Generate charts
        charts.append(self.generate_severity_distribution(timeline_data))
        charts.append(self.generate_timeline_heatmap(timeline_data))
        charts.append(self.generate_source_analysis(timeline_data))
        charts.append(self.generate_event_timeline(timeline_data))
        
        if iocs:
            charts.append(self.generate_ioc_analysis(iocs))
        
        if mitre_mappings:
            charts.append(self.generate_mitre_analysis(mitre_mappings))
        
        if attack_chain:
            charts.append(self.generate_attack_chain_visualization(attack_chain))
        
        return charts

    def generate_event_type_analysis(self, timeline_data: List[Dict[str, Any]]) -> go.Figure:
        """Generate event type analysis chart."""
        if not timeline_data:
            return self._create_empty_chart("No timeline data available")
        
        # Count events by type
        event_type_counts = {}
        for event in timeline_data:
            event_type = event.get('event_type', 'unknown')
            event_type_counts[event_type] = event_type_counts.get(event_type, 0) + 1
        
        if not event_type_counts:
            return self._create_empty_chart("No event type data found")
        
        # Create horizontal bar chart
        event_types = list(event_type_counts.keys())
        counts = list(event_type_counts.values())
        
        fig = go.Figure(data=[go.Bar(
            y=event_types,
            x=counts,
            orientation='h',
            marker_color='lightcoral',
            text=counts,
            textposition='auto'
        )])
        
        fig.update_layout(
            title="Events by Type",
            xaxis_title="Event Count",
            yaxis_title="Event Type",
            height=400,
            margin=dict(t=50, b=50, l=150, r=50)  # Extra left margin for long event type names
        )
        
        return fig

    def generate_confidence_analysis(self, report: ForensicReport) -> go.Figure:
        """Generate confidence analysis across different analysis components."""
        confidence_data = {}
        
        # Get confidence from different sources
        if getattr(report, 'analysis', None) and getattr(report.analysis, 'initial_analysis', None):
            initial_conf = float(getattr(report.analysis.initial_analysis, 'confidence', 0.0))
            confidence_data['Initial Analysis'] = initial_conf
        
        # Get IOC confidence scores
        iocs = getattr(getattr(report.analysis, 'initial_analysis', {}), 'iocs', []) if getattr(report, 'analysis', None) else []
        if iocs:
            ioc_confidences = [float(getattr(ioc, 'confidence', 0.0)) for ioc in iocs if getattr(ioc, 'confidence', None) is not None]
            if ioc_confidences:
                confidence_data['Average IOC Confidence'] = sum(ioc_confidences) / len(ioc_confidences)
        
        # Get MITRE confidence scores
        mitre_mappings = getattr(getattr(report.analysis, 'initial_analysis', {}), 'mitre_mappings', []) if getattr(report, 'analysis', None) else []
        if mitre_mappings:
            mitre_confidences = [float(mapping.confidence) for mapping in mitre_mappings if getattr(mapping, 'confidence', None) is not None]
            if mitre_confidences:
                confidence_data['Average MITRE Confidence'] = sum(mitre_confidences) / len(mitre_confidences)
        
        if not confidence_data:
            return self._create_empty_chart("No confidence data available")
        
        # Create gauge chart
        fig = go.Figure()
        
        for i, (component, confidence) in enumerate(confidence_data.items()):
            fig.add_trace(go.Indicator(
                mode="gauge+number+delta",
                value=confidence * 100,  # Convert to percentage
                domain={'x': [0, 1], 'y': [0, 1]},
                title={'text': component},
                delta={'reference': 50},
                gauge={
                    'axis': {'range': [None, 100]},
                    'bar': {'color': "darkblue"},
                    'steps': [
                        {'range': [0, 50], 'color': "lightgray"},
                        {'range': [50, 80], 'color': "yellow"},
                        {'range': [80, 100], 'color': "green"}
                    ],
                    'threshold': {
                        'line': {'color': "red", 'width': 4},
                        'thickness': 0.75,
                        'value': 90
                    }
                }
            ))
        
        fig.update_layout(
            title="Analysis Confidence Overview",
            height=300 * len(confidence_data),
            margin=dict(t=50, b=50, l=50, r=50)
        )
        
        return fig

    def _create_empty_chart(self, message: str) -> go.Figure:
        """Create an empty chart with a message."""
        fig = go.Figure()
        fig.add_annotation(
            text=message,
            xref="paper", yref="paper",
            x=0.5, y=0.5,
            showarrow=False,
            font=dict(size=16, color="gray")
        )
        fig.update_layout(
            xaxis=dict(visible=False),
            yaxis=dict(visible=False),
            height=300,
            margin=dict(t=50, b=50, l=50, r=50)
        )
        return fig


# -----------------------------------------------------------------------------
# Streamlit App
# -----------------------------------------------------------------------------
class ForensicDashboard:
    """Main Streamlit dashboard for forensic analysis."""

    def __init__(self):
        self.orchestrator = ForensicOrchestrator()
        self.report_generator = ReportGenerator()
        self.chart_generator = ChartGenerator()

        # Configure Streamlit page
        st.set_page_config(
            page_title="Forensic AI Log Analyzer",
            layout="wide",
            initial_sidebar_state="expanded",
        )

    def run(self):
        """Main dashboard application."""
        st.title("Forensic AI Log Analyzer")
        st.markdown("---")

        # Sidebar
        self._render_sidebar()

        # Tabs
        tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
            "Upload & Analyze",
            "Analysis Results",
            "Timeline View",
            "Dashboard",
            "Reports",
            "Semantic Search",
            "Saved Results",
            "File History",
        ])

        with tab1:
            self._render_upload_tab()
        with tab2:
            self._render_analysis_tab()
        with tab3:
            self._render_timeline_tab()
        with tab4:
            self._render_dashboard_tab()
        with tab5:
            self._render_reports_tab()
        with tab6:
            self._render_semantic_search_tab()
        with tab7:
            self._render_saved_results_tab()
        with tab8:
            self._render_file_history_tab()

    # ---------------------------------------------
    # Sidebar
    # ---------------------------------------------
    def _render_sidebar(self):
        st.sidebar.header("Configuration")

        # LLM provider controls (runtime override)
        st.sidebar.subheader("LLM Provider")
        current_provider = getattr(Config, "LLM_PROVIDER", "ibm")
        selected_provider = st.sidebar.selectbox("Provider", ["ibm", "groq"], index=(0 if current_provider != "groq" else 1))
        if selected_provider != current_provider:
            Config.LLM_PROVIDER = selected_provider
        # if selected_provider == "groq":
        #     groq_key_input = st.sidebar.text_input("Groq API Key", value=(Config.GROQ_API_KEY or ""), type="password")
        #     if groq_key_input and groq_key_input != Config.GROQ_API_KEY:
        #         Config.GROQ_API_KEY = groq_key_input

        # Configuration status
        config_status = Config.validate_config()
        st.sidebar.subheader("Service Status")
        for service, status in config_status.items():
            service_name = service.replace("_", " ").title()
            st.sidebar.text(f"{service_name}: {'OK' if status else 'Not Configured'}")

        # Explicit LLM provider status
        provider = getattr(Config, "LLM_PROVIDER", "ibm")
        # Cache analyzer probe to avoid re-instantiation overhead
        if "_analyzer_probe" not in st.session_state:
            st.session_state["_analyzer_probe"] = GraniteAnalyzer()
        analyzer_probe = st.session_state["_analyzer_probe"]
        if provider == "groq":
            llm_ready = bool(getattr(analyzer_probe, "groq_client", None))
        else:
            llm_ready = bool(getattr(analyzer_probe, "model", None))
        st.sidebar.text(f"LLM Provider: {provider.upper()} - {'Ready' if llm_ready else 'Not Ready'}")
        if provider == "groq" and not llm_ready:
            st.sidebar.warning("Groq selected but not initialized. Provide API key above and try again.")

        # Analysis settings
        st.sidebar.subheader("Analysis Settings")
        st.sidebar.slider("Max Log Entries", 100, 50000, Config.MAX_LOG_ENTRIES, key="max_entries")
        st.sidebar.selectbox("Analysis Depth", ["Quick", "Standard", "Deep"], index=1, key="analysis_depth")
        
        # Visualization settings
        st.sidebar.subheader("Visualization Settings")
        show_charts = st.sidebar.checkbox("Enable charts and visualizations", value=True, key="enable_charts")
        if show_charts:
            chart_quality = st.sidebar.selectbox(
                "Chart Quality", 
                ["Standard", "High Quality"], 
                index=0,
                help="Higher quality charts may take longer to render"
            )
            st.sidebar.checkbox(
                "Interactive charts", 
                value=True, 
                key="interactive_charts",
                help="Enable hover effects and zoom in charts"
            )

        # About
        st.sidebar.subheader("About")
        st.sidebar.info(
            "Forensic AI Log Analyzer v2.0\n\n"
            "Enhanced AI-powered tool for automated security log analysis and incident investigation with comprehensive visualizations and reporting."
        )
        
        # New Features
        st.sidebar.subheader("✨ New Features")
        st.sidebar.success(
            "• Enhanced Charts & Visualizations\n"
            "• Comprehensive Dashboard\n"
            "• Professional HTML Reports\n"
            "• Interactive Timeline Analysis\n"
            "• Confidence Analysis\n"
            "• Attack Chain Visualization"
        )

    # ---------------------------------------------
    # Upload & Analyze
    # ---------------------------------------------
    def _render_upload_tab(self):
        st.header("Upload Log Files")

        # S3 Status and Configuration
        s3_available = is_s3_available()
        if s3_available:
            st.success("✅ AWS S3 is configured and available")
        else:
            st.info("ℹ️ AWS S3 is not configured. Files will be processed locally only.")
            st.info("To enable S3, set AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, and AWS_S3_BUCKET environment variables.")

        # S3 File Selection (if available)
        if s3_available:
            st.subheader("📁 Select Files from S3")
            
            s3_utils = get_s3_utils()
            s3_files = s3_utils.list_files()
            
            if s3_files:
                selected_s3_file = st.selectbox(
                    "Choose a file from S3",
                    options=[""] + s3_files,
                    format_func=lambda x: "Select a file..." if x == "" else x,
                    help="Select a file from your S3 bucket to analyze"
                )
                
                if selected_s3_file:
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.info(f"Selected: {selected_s3_file}")
                    with col2:
                        if st.button("Download & Analyze", key="s3_download_analyze"):
                            with st.spinner(f"Downloading {selected_s3_file} from S3..."):
                                try:
                                    downloaded_path = s3_utils.download_file(key=selected_s3_file)
                                    if downloaded_path:
                                        st.success(f"Downloaded to: {downloaded_path}")
                                        # Add to temp_files for analysis
                                        if 'temp_files' not in st.session_state:
                                            st.session_state.temp_files = []
                                        st.session_state.temp_files.append(downloaded_path)
                                        
                                        # Add file info
                                        if 'file_info' not in st.session_state:
                                            st.session_state.file_info = []
                                        file_info_item = {
                                            "Name": selected_s3_file,
                                            "Size": "Unknown",
                                            "Type": "S3 File",
                                            "Source": "S3 Bucket"
                                        }
                                        st.session_state.file_info.append(file_info_item)
                                        st.rerun()
                                    else:
                                        st.error("Failed to download file from S3")
                                except Exception as e:
                                    st.error(f"Error downloading file: {e}")
                                    logger.exception("S3 download error")
            else:
                st.info("No files found in S3 bucket")

        # File upload section
        st.subheader("📤 Upload Files")
        
        # ZIP file upload
        zip_file = st.file_uploader(
            "Upload ZIP file containing log files",
            type=["zip"],
            help="Upload a ZIP file containing multiple log files (network, system, process, services, etc.)",
            key="zip_uploader"
        )
        
        # Individual file upload
        uploaded_files = st.file_uploader(
            "Or upload individual log files",
            accept_multiple_files=True,
            type=["csv", "json", "evtx", "log", "txt"],
            help="Supported formats: CSV, JSON, EVTX, LOG, TXT",
            key="file_uploader"
        )

        # Process uploaded files
        temp_files: List[str] = []
        file_info = []
        zip_processor = None
        zip_results = None
        
        # Add S3 files from session state
        if 'temp_files' in st.session_state and st.session_state.temp_files:
            temp_files.extend(st.session_state.temp_files)
        if 'file_info' in st.session_state and st.session_state.file_info:
            file_info.extend(st.session_state.file_info)
        
        # Handle ZIP file upload
        if zip_file:
            st.subheader("ZIP File Processing")
            
            try:
                # Save ZIP file to temporary location
                with tempfile.NamedTemporaryFile(delete=False, suffix=".zip") as tmp_zip:
                    tmp_zip.write(zip_file.getbuffer())
                    zip_path = tmp_zip.name
                
                # Process ZIP file
                from zip_processor import ZIPProcessor
                zip_processor = ZIPProcessor(max_file_size_mb=100, max_total_size_mb=500)
                zip_results = zip_processor.extract_and_categorize(zip_path)
                
                if zip_results["success"]:
                    st.success(f"Successfully extracted {zip_results['file_count']} files from ZIP")
                    
                    # Display ZIP summary
                    with st.expander("ZIP Extraction Summary", expanded=True):
                        st.text(zip_processor.generate_summary_report(zip_results))
                    
                    # Display categorized files
                    st.subheader("Extracted Files by Category")
                    
                    for category, files in zip_results["categorized_files"].items():
                        with st.expander(f"{category.title()} ({len(files)} files)", expanded=True):
                            category_df = pd.DataFrame([
                                {
                                    "Name": f["name"],
                                    "Size": f"{f['size_mb']:.2f} MB",
                                    "Path": f["relative_path"]
                                }
                                for f in files
                            ])
                            st.dataframe(category_df, width='stretch')
                    
                    # Get readable files for analysis
                    readable_files = zip_processor.get_readable_files(zip_results["extracted_files"])
                    temp_files.extend(readable_files)
                    
                    # Add file info for display
                    for file_info_item in zip_results["extracted_files"]:
                        file_info.append({
                            "Name": file_info_item["name"],
                            "Size": f"{file_info_item['size_mb']:.2f} MB",
                            "Type": file_info_item["category"].title(),
                            "Source": "ZIP Archive"
                        })
                    
                    # Upload ZIP contents to S3 if requested
                    if s3_available and st.checkbox("Upload ZIP contents to S3", value=False, key="zip_s3_upload"):
                        s3_utils = get_s3_utils()
                        uploaded_count = 0
                        for file_path in readable_files:
                            try:
                                file_name = Path(file_path).name
                                s3_key = f"zip_extracts/{datetime.now().strftime('%Y%m%d_%H%M%S')}/{file_name}"
                                if s3_utils.upload_file(file_path, key=s3_key):
                                    uploaded_count += 1
                            except Exception as e:
                                logger.warning(f"Failed to upload {file_path} to S3: {e}")
                        
                        if uploaded_count > 0:
                            st.success(f"✅ Uploaded {uploaded_count} files from ZIP to S3")
                    
                    # Clean up temporary ZIP file
                    os.unlink(zip_path)
                    
                else:
                    st.error(f"Failed to extract ZIP file: {zip_results['error']}")
                    if zip_results["temp_dir"]:
                        zip_processor.cleanup_temp_dir(zip_results["temp_dir"])
                    
            except Exception as e:
                st.error(f"Error processing ZIP file: {str(e)}")
                logger.exception("ZIP processing error")
        
        # Handle individual file uploads
        if uploaded_files:
            st.subheader("Individual Files")
            
            # S3 upload option
            if s3_available:
                upload_to_s3 = st.checkbox(
                    "Upload files to S3 after processing",
                    value=False,
                    help="Upload processed files to your S3 bucket for future access"
                )
            else:
                upload_to_s3 = False
            
            for uploaded_file in uploaded_files:
                with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{uploaded_file.name}") as tmp_file:
                    tmp_file.write(uploaded_file.getbuffer())
                    temp_files.append(tmp_file.name)
                    
                    # Upload to S3 if requested
                    if upload_to_s3 and s3_available:
                        try:
                            s3_utils = get_s3_utils()
                            s3_key = f"uploads/{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uploaded_file.name}"
                            if s3_utils.upload_file(tmp_file.name, key=s3_key):
                                st.success(f"✅ Uploaded {uploaded_file.name} to S3 as {s3_key}")
                            else:
                                st.warning(f"⚠️ Failed to upload {uploaded_file.name} to S3")
                        except Exception as e:
                            st.warning(f"⚠️ S3 upload error for {uploaded_file.name}: {e}")
                            logger.exception(f"S3 upload error for {uploaded_file.name}")

                file_info.append(
                    {
                        "Name": uploaded_file.name,
                        "Size": f"{len(uploaded_file.getbuffer()) / 1024:.1f} KB",
                        "Type": uploaded_file.type or "Unknown",
                        "Source": "Direct Upload"
                    }
                )
        
        # Display all files if any were uploaded
        if file_info:
            st.subheader("All Files for Analysis")
            df_files = pd.DataFrame(file_info)
            st.dataframe(df_files, width='stretch')
            
            # Check for EVTX files and show warning
            evtx_files = [f for f in file_info if f["Type"].lower() == "evtx"]
            if evtx_files and not skip_evtx:
                st.warning("**EVTX files detected**: Windows Event Log files (.evtx) can cause parsing delays or timeouts. If analysis gets stuck, try enabling 'Skip EVTX Files' option above.")
            
            # Show file statistics
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Total Files", len(file_info))
            with col2:
                total_size = sum(float(f["Size"].split()[0]) for f in file_info if "MB" in f["Size"])
                st.metric("Total Size", f"{total_size:.2f} MB")
            with col3:
                categories = set(f["Type"] for f in file_info)
                st.metric("Categories", len(categories))

            # Analysis controls
            st.markdown("---")
            st.subheader("Analysis Configuration")
            
            col1, col2, col3, col4, col5, col6 = st.columns([2, 1, 1, 1, 1, 1])
            with col1:
                case_id = st.text_input("Case ID (optional)", placeholder="CASE_2025_001")
            with col2:
                skip_chromadb = st.checkbox("Skip ChromaDB", value=False, help="Skip storing logs for semantic search to speed up analysis")
            with col3:
                skip_evtx = st.checkbox("Skip EVTX Files", value=False, help="Skip Windows Event Log files to avoid parsing issues. Enable if analysis gets stuck on .evtx files.")
            with col4:
                skip_rag = st.checkbox("Skip RAG Enrichment", value=False, help="Skip RAG enrichment to avoid hanging. Enable if analysis gets stuck at 'RAG enrichment completed successfully'.")
            with col5:
                processing_mode = st.selectbox(
                    "Processing Mode",
                    ["High Performance", "Standard"],
                    help="High Performance: Parallel processing for large datasets. Standard: Sequential processing for smaller datasets."
                )
            with col6:
                st.write("")
                analyze_button = st.button("Start Analysis", type="primary")

            if analyze_button:
                with st.spinner("Running forensic analysis..."):
                    try:
                        progress_bar = st.progress(0)
                        status_text = st.empty()

                        status_text.text("Parsing log files and storing in DynamoDB...")
                        progress_bar.progress(25)

                        # Store logs in ChromaDB for semantic search
                        try:
                            from chromadb_logs import ingest_logs, ingest_logs_high_performance
                            
                            # Test ChromaDB connection first
                            status_text.text("Testing ChromaDB connection...")
                            try:
                                from chromadb_logs import get_log_store
                                log_store = get_log_store()
                                stats = log_store.get_collection_stats()
                                st.success(f"ChromaDB connection verified - {stats.get('total_logs', 0)} existing logs")
                            except Exception as e:
                                st.error(f"ChromaDB connection failed: {str(e)[:100]}...")
                                st.warning("Please check your ChromaDB configuration")
                                st.session_state["stored_logs_count"] = 0
                                raise e
                            
                            # Read and parse log files
                            status_text.text("Reading and parsing log files...")
                            log_entries = []
                            for temp_file in temp_files:
                                try:
                                    with open(temp_file, 'r', encoding='utf-8', errors='ignore') as f:
                                        for line_num, line in enumerate(f, 1):
                                            line = line.strip()
                                            if line and len(line) > 10:  # Skip empty/short lines
                                                log_entries.append(line)
                                            
                                            # Limit to prevent overwhelming the system
                                            if len(log_entries) >= 1000:
                                                break
                                    
                                    if len(log_entries) >= 1000:
                                        break
                                except Exception as e:
                                    logger.warning(f"Could not read file {temp_file}: {e}")
                            
                            st.info(f"Parsed {len(log_entries)} valid log entries")
                            
                            if log_entries and not skip_chromadb:
                                # Store in ChromaDB with embeddings
                                status_text.text("Generating embeddings and storing in ChromaDB...")
                                progress_bar.progress(35)
                                
                                try:
                                    # Store logs using selected processing mode
                                    if processing_mode == "High Performance":
                                        # High-performance mode: Parallel processing for large datasets
                                        max_workers = min(8, len(log_entries) // 100 + 1)  # Scale workers based on dataset size
                                        stored_items = ingest_logs_high_performance(log_entries, max_workers=max_workers)
                                        logger.info(f"Using high-performance mode with {max_workers} workers")
                                    else:
                                        # Standard mode: Sequential processing for smaller datasets
                                        stored_items = ingest_logs(log_entries)
                                        logger.info("Using standard sequential processing mode")
                                    
                                    if isinstance(stored_items, dict) and stored_items.get("stored_items"):
                                        # High-performance function returned detailed results
                                        successful_count = stored_items["successful"]
                                        failed_count = stored_items["failed"]
                                        st.session_state["stored_logs_count"] = successful_count
                                        
                                        if failed_count > 0:
                                            st.warning(f"Stored {successful_count} logs in ChromaDB, {failed_count} failed")
                                        else:
                                            st.success(f"Stored {successful_count} logs in ChromaDB")
                                        
                                        logger.info(f"Successfully stored {successful_count} log entries in ChromaDB using high-performance batching")
                                        progress_bar.progress(50)
                                    elif stored_items:
                                        # Regular function returned simple list
                                        st.session_state["stored_logs_count"] = len(stored_items)
                                        logger.info(f"Successfully stored {len(stored_items)} log entries in ChromaDB")
                                        st.success(f"Stored {len(stored_items)} logs in ChromaDB")
                                        progress_bar.progress(50)
                                    else:
                                        st.warning("No logs were stored in ChromaDB")
                                        st.session_state["stored_logs_count"] = 0
                                        
                                except Exception as e:
                                    logger.error(f"Failed to store logs in ChromaDB: {e}")
                                    st.error(f"Failed to store logs in ChromaDB: {str(e)}")
                                    st.session_state["stored_logs_count"] = 0
                                    # Continue with analysis even if ChromaDB fails
                            elif skip_chromadb:
                                st.info("Skipping ChromaDB storage as requested")
                                st.session_state["stored_logs_count"] = 0
                            else:
                                st.warning("No valid log entries found to store")
                                st.session_state["stored_logs_count"] = 0
                                
                        except Exception as e:
                            logger.error(f"Failed to store logs in ChromaDB: {e}")
                            st.warning("Logs could not be stored in ChromaDB for semantic search")
                            st.session_state["stored_logs_count"] = 0

                        # Track file history
                        all_uploaded_files = []
                        if zip_file:
                            all_uploaded_files.append(zip_file)
                        if uploaded_files:
                            all_uploaded_files.extend(uploaded_files)
                        
                        self._update_file_history(all_uploaded_files, case_id, len(log_entries) if 'log_entries' in locals() else 0)

                        status_text.text("Running forensic investigation...")
                        progress_bar.progress(50)

                        # Run investigation with EVTX skip option
                        if skip_evtx:
                            # Create a new orchestrator with EVTX files skipped
                            from agents import ForensicOrchestrator
                            evtx_skip_orchestrator = ForensicOrchestrator(skip_evtx_files=True)
                            report = evtx_skip_orchestrator.investigate(temp_files, case_id or None, skip_rag)
                        else:
                            # Use the default orchestrator
                            report = self.orchestrator.investigate(temp_files, case_id or None, skip_rag)

                        status_text.text("Running AI analysis...")
                        progress_bar.progress(60)

                        # Generate report (timeout handling is now in the RAG pipeline)
                        status_text.text("Generating report...")
                        progress_bar.progress(100)

                        st.session_state["forensic_report"] = report
                        st.session_state["analysis_complete"] = True

                        # Save analysis result to database
                        try:
                            from results_storage import get_results_storage
                            storage = get_results_storage()
                            
                            # Prepare analysis configuration
                            analysis_config = {
                                "processing_mode": processing_mode,
                                "skip_chromadb": skip_chromadb,
                                "max_entries": st.session_state.get("max_entries", Config.MAX_LOG_ENTRIES),
                                "analysis_depth": st.session_state.get("analysis_depth", "Standard")
                            }
                            
                            # Save the result
                            analysis_id = storage.save_analysis_result(
                                case_id=case_id or f"auto_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                                case_name=case_id or "Auto-generated Case",
                                analysis_result=report,
                                file_paths=temp_files,
                                zip_file_path=zip_file.name if zip_file else None,
                                analysis_config=analysis_config
                            )
                            
                            st.session_state["saved_analysis_id"] = analysis_id
                            st.success(f"Analysis saved with ID: {analysis_id}")
                            
                        except Exception as e:
                            logger.error(f"Failed to save analysis result: {e}")
                            st.warning("Analysis completed but could not be saved to database")

                        # Cleanup temp files
                        for p in temp_files:
                            try:
                                Path(p).unlink(missing_ok=True)
                            except Exception:
                                pass
                        
                        # Cleanup ZIP extraction directory if it exists
                        if zip_processor and zip_results and zip_results.get("temp_dir"):
                            zip_processor.cleanup_temp_dir(zip_results["temp_dir"])
                        
                        # Cleanup S3 session state
                        if 'temp_files' in st.session_state:
                            del st.session_state.temp_files
                        if 'file_info' in st.session_state:
                            del st.session_state.file_info
                        
                        # Cleanup S3 session state
                        if 'temp_files' in st.session_state:
                            del st.session_state.temp_files
                        if 'file_info' in st.session_state:
                            del st.session_state.file_info

                        status_text.text("Analysis complete.")
                        success_msg = f"Forensic analysis completed. Case ID: {getattr(report, 'case_id', 'N/A')}"
                        if st.session_state.get("stored_logs_count"):
                            success_msg += f" | {st.session_state['stored_logs_count']} logs stored in ChromaDB"
                        st.success(success_msg)
                    except Exception as e:  # pragma: no cover
                        st.error(f"Analysis failed: {e}")
                        logger.exception("Dashboard analysis error")
        else:
            st.info("Upload log files or a ZIP archive to begin analysis")
            
            # Show supported formats
            st.markdown("---")
            st.subheader("Supported Formats")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Individual Files:**")
                st.markdown("• CSV files (.csv)")
                st.markdown("• JSON files (.json)")
                st.markdown("• Windows Event Logs (.evtx)")
                st.markdown("• Log files (.log)")
                st.markdown("• Text files (.txt)")
            
            with col2:
                st.markdown("**ZIP Archives:**")
                st.markdown("• Network logs (pcap, netflow)")
                st.markdown("• System logs (syslog, events)")
                st.markdown("• Process lists (tasklist, ps)")
                st.markdown("• Service logs (daemon, init)")
                st.markdown("• Authentication logs (auth, login)")
                st.markdown("• Any combination of the above")

    # ---------------------------------------------
    # Analysis Results
    # ---------------------------------------------
    def _render_analysis_tab(self):
        # Check if we're viewing a saved analysis
        if st.session_state.get("view_analysis_id"):
            self._render_saved_analysis(st.session_state["view_analysis_id"])
            return
        
        if not st.session_state.get("analysis_complete", False):
            st.info("Run analysis first to view results")
            return

        report: ForensicReport = st.session_state.get("forensic_report")
        if not report:
            st.error("No analysis results available")
            return

        st.header("Analysis Results")

        # Executive Summary
        st.subheader("Executive Summary")
        st.info(str(getattr(report, "executive_summary", "No summary provided.")))

        # Key metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Case ID", str(getattr(report, "case_id", "N/A")))
        with col2:
            st.metric("Total Events", f"{getattr(report, 'total_events', 0):,}")
        with col3:
            sev = (
                str(getattr(getattr(report.analysis.initial_analysis, "severity", ""), "value", "")).upper()
                if getattr(report, "analysis", None) and getattr(report.analysis, "initial_analysis", None)
                else "N/A"
            )
            st.metric("Severity", sev)
        with col4:
            conf = (
                f"{float(getattr(report.analysis.initial_analysis, 'confidence', 0.0)):.1%}"
                if getattr(report, "analysis", None) and getattr(report.analysis, "initial_analysis", None)
                else "N/A"
            )
            st.metric("Confidence", conf)

        # IOCs
        st.subheader("Indicators of Compromise")
        iocs = (
            getattr(getattr(report.analysis, "initial_analysis", {}), "iocs", [])
            if getattr(report, "analysis", None)
            else []
        )
        if iocs:
            ioc_data = []
            for ioc in iocs:
                ioc_data.append(
                    {
                        "Type": str(getattr(ioc, "type", "")).upper(),
                        "Value": str(getattr(ioc, "value", "")),
                        "Confidence": (
                            f"{float(getattr(ioc, 'confidence', 0.0)):.1%}" if getattr(ioc, "confidence", None) is not None else "N/A"
                        ),
                        "Context": str(getattr(ioc, "context", "")) or "N/A",
                    }
                )
            df_iocs = pd.DataFrame(ioc_data)
            st.dataframe(df_iocs, width='stretch')

            # Enhanced IOC visualization is now in the Dashboard tab
        else:
            st.info("No IOCs identified")

        # MITRE ATT&CK
        st.subheader("MITRE ATT&CK Mapping")
        mitre = (
            getattr(getattr(report.analysis, "initial_analysis", {}), "mitre_mappings", [])
            if getattr(report, "analysis", None)
            else []
        )
        if mitre:
            mitre_rows = []
            for m in mitre:
                mitre_rows.append(
                    {
                        "Tactic ID": getattr(m, "tactic_id", ""),
                        "Tactic": getattr(m, "tactic_name", ""),
                        "Technique ID": getattr(m, "technique_id", "N/A"),
                        "Technique": getattr(m, "technique_name", "N/A"),
                        "Confidence": (
                            f"{float(getattr(m, 'confidence', 0.0)):.1%}"
                            if getattr(m, "confidence", None) is not None
                            else "N/A"
                        ),
                    }
                )
            df_mitre = pd.DataFrame(mitre_rows)
            st.dataframe(df_mitre, width='stretch')
        else:
            st.info("No MITRE ATT&CK mappings identified")

        # Attack chain visualization
        st.subheader("Attack Chain")
        chain = getattr(getattr(report, "analysis", {}), "attack_chain", None)
        if chain:
            self._render_attack_chain(chain)
        else:
            st.info("Attack chain could not be reconstructed")

        # Recommendations
        st.subheader("Recommendations")
        recs = getattr(getattr(report, "analysis", {}), "recommendations", [])
        if recs:
            for i, r in enumerate(recs, 1):
                st.write(f"{i}. {r}")
        else:
            st.info("No specific recommendations generated")

        # Enhanced Charts and Visualizations
        st.markdown("---")
        st.subheader("📊 Analysis Visualizations")
        
        if st.session_state.get("enable_charts"):
            # Generate comprehensive charts
            charts = self.chart_generator.generate_comprehensive_dashboard(report)
            
            # Display charts in organized sections
            if charts:
                # Severity and Timeline Analysis
                col1, col2 = st.columns(2)
                with col1:
                    st.plotly_chart(charts[0], use_container_width=True)  # Severity distribution
                with col2:
                    if len(charts) > 1:
                        st.plotly_chart(charts[1], use_container_width=True)  # Timeline heatmap
                
                # Source Analysis and Event Timeline
                if len(charts) > 2:
                    st.plotly_chart(charts[2], use_container_width=True)  # Source analysis
                
                if len(charts) > 3:
                    st.plotly_chart(charts[3], use_container_width=True)  # Event timeline
                
                # IOC and MITRE Analysis
                if len(charts) > 4:
                    col1, col2 = st.columns(2)
                    with col1:
                        st.plotly_chart(charts[4], use_container_width=True)  # IOC analysis
                    with col2:
                        if len(charts) > 5:
                            st.plotly_chart(charts[5], use_container_width=True)  # MITRE analysis
                
                # Attack Chain Visualization
                if len(charts) > 6:
                    st.plotly_chart(charts[6], use_container_width=True)  # Attack chain
        else:
            st.info("Enable charts in the sidebar to view comprehensive visualizations")
            
            # Show chart preview
            with st.expander("Chart Preview"):
                st.markdown("""
                **Available Charts:**
                - **Event Severity Distribution**: Pie chart showing distribution of event severities
                - **Timeline Heatmap**: Heatmap showing event density over time
                - **Source Analysis**: Bar chart showing events by source
                - **Interactive Timeline**: Scatter plot of events over time
                - **IOC Analysis**: Distribution of IOC types and confidence scores
                - **MITRE ATT&CK Analysis**: Tactics and techniques distribution
                - **Attack Chain Flow**: Sankey diagram showing attack progression
                """)

    # ---------------------------------------------
    # Dashboard
    # ---------------------------------------------
    def _render_dashboard_tab(self):
        """Render comprehensive dashboard with all visualizations."""
        if not st.session_state.get("analysis_complete", False):
            st.info("Run analysis first to view dashboard")
            return

        report: ForensicReport = st.session_state.get("forensic_report")
        if not report:
            st.error("No analysis results available")
            return

        st.header("📊 Comprehensive Analysis Dashboard")
        st.markdown("---")

        # Key metrics overview
        st.subheader("🎯 Key Metrics")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Case ID", str(getattr(report, "case_id", "N/A")))
        with col2:
            st.metric("Total Events", f"{getattr(report, 'total_events', 0):,}")
        with col3:
            sev = (
                str(getattr(getattr(report.analysis.initial_analysis, "severity", ""), "value", "")).upper()
                if getattr(report, "analysis", None) and getattr(report.analysis, "initial_analysis", None)
                else "N/A"
            )
            st.metric("Severity", sev)
        with col4:
            conf = (
                f"{float(getattr(report.analysis.initial_analysis, 'confidence', 0.0)):.1%}"
                if getattr(report, "analysis", None) and getattr(report.analysis, "initial_analysis", None)
                else "N/A"
            )
            st.metric("Confidence", conf)

        # Executive Summary
        st.subheader("📋 Executive Summary")
        st.info(str(getattr(report, "executive_summary", "No summary provided.")))

        # Comprehensive Charts
        st.subheader("📈 Analysis Visualizations")
        
        if st.session_state.get("enable_charts"):
            # Generate all charts
            charts = self.chart_generator.generate_comprehensive_dashboard(report)
            
            if charts:
                # Overview charts
                st.markdown("### Event Overview")
                col1, col2 = st.columns(2)
                with col1:
                    st.plotly_chart(charts[0], use_container_width=True)  # Severity distribution
                with col2:
                    if len(charts) > 1:
                        st.plotly_chart(charts[1], use_container_width=True)  # Timeline heatmap
                
                # Source and Timeline Analysis
                st.markdown("### Source and Timeline Analysis")
                if len(charts) > 2:
                    st.plotly_chart(charts[2], use_container_width=True)  # Source analysis
                
                if len(charts) > 3:
                    st.plotly_chart(charts[3], use_container_width=True)  # Event timeline
                
                # Threat Intelligence
                st.markdown("### Threat Intelligence Analysis")
                if len(charts) > 4:
                    col1, col2 = st.columns(2)
                    with col1:
                        st.plotly_chart(charts[4], use_container_width=True)  # IOC analysis
                    with col2:
                        if len(charts) > 5:
                            st.plotly_chart(charts[5], use_container_width=True)  # MITRE analysis
                
                # Attack Chain
                if len(charts) > 6:
                    st.markdown("### Attack Chain Visualization")
                    st.plotly_chart(charts[6], use_container_width=True)  # Attack chain
                
                # Additional Analysis
                st.markdown("### Additional Analysis")
                col1, col2 = st.columns(2)
                with col1:
                    event_type_fig = self.chart_generator.generate_event_type_analysis(timeline_data)
                    st.plotly_chart(event_type_fig, use_container_width=True)
                
                with col2:
                    confidence_fig = self.chart_generator.generate_confidence_analysis(report)
                    st.plotly_chart(confidence_fig, use_container_width=True)
        else:
            st.info("Enable charts in the sidebar to view comprehensive visualizations")
            
            # Show what's available
            with st.expander("Available Dashboard Features"):
                st.markdown("""
                **Dashboard Components:**
                - **Event Severity Distribution**: Visual breakdown of security event severities
                - **Timeline Heatmap**: Event density analysis over time
                - **Source Analysis**: Events categorized by source systems
                - **Interactive Timeline**: Detailed event timeline with filtering
                - **IOC Analysis**: Distribution and confidence of identified IOCs
                - **MITRE ATT&CK Mapping**: Tactics and techniques visualization
                - **Attack Chain Flow**: Visual representation of attack progression
                """)

        # Quick Actions
        st.markdown("---")
        st.subheader("⚡ Quick Actions")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📥 Download Report", key="dashboard_download"):
                st.info("Use the Reports tab to download comprehensive reports")
        
        with col2:
            if st.button("🔍 View Timeline", key="dashboard_timeline"):
                st.info("Switch to Timeline View tab for detailed event analysis")
        
        with col3:
            if st.button("📊 Export Charts", key="dashboard_export"):
                st.info("Chart export functionality coming soon!")

    # ---------------------------------------------
    # Timeline View
    # ---------------------------------------------
    def _render_timeline_tab(self):
        if not st.session_state.get("analysis_complete", False):
            st.info("Run analysis first to view timeline")
            return

        report: ForensicReport = st.session_state.get("forensic_report")
        if not report or not getattr(report, "timeline", None):
            st.error("No timeline data available")
            return

        st.header("Event Timeline")

        # Controls
        col1, col2 = st.columns([3, 1])
        with col1:
            times = [
                datetime.fromisoformat(e.get("timestamp", "").replace("Z", "+00:00"))
                for e in report.timeline
                if e.get("timestamp")
            ]
            if not times:
                st.info("Timeline has no timestamps to plot")
                return
            min_time, max_time = min(times), max(times)
            time_range = st.slider(
                "Time Range",
                min_value=min_time,
                max_value=max_time,
                value=(min_time, max_time),
                format="YYYY-MM-DD HH:mm",
            )
        with col2:
            severity_filter = st.multiselect(
                "Severity Filter",
                ["critical", "high", "medium", "low", "info"],
                default=["critical", "high", "medium", "low", "info"],
            )

        # Filter
        filtered = []
        for ev in report.timeline:
            try:
                ev_time = datetime.fromisoformat(ev["timestamp"].replace("Z", "+00:00"))
            except Exception:
                continue
            if time_range[0] <= ev_time <= time_range[1] and ev.get("severity") in severity_filter:
                filtered.append(ev)

        if filtered:
            df = pd.DataFrame(filtered)
            df["timestamp"] = pd.to_datetime(df["timestamp"])
            if "sequence" not in df.columns:
                df["sequence"] = range(1, len(df) + 1)
            if "source" not in df.columns:
                df["source"] = "unknown"

            if st.session_state.get("enable_charts"):
                # Enhanced timeline visualization
                fig = self.chart_generator.generate_event_timeline(filtered)
                st.plotly_chart(fig, use_container_width=True)
                
                # Additional timeline analysis
                col1, col2 = st.columns(2)
                with col1:
                    severity_fig = self.chart_generator.generate_severity_distribution(filtered)
                    st.plotly_chart(severity_fig, use_container_width=True)
                
                with col2:
                    source_fig = self.chart_generator.generate_source_analysis(filtered)
                    st.plotly_chart(source_fig, use_container_width=True)
                
                # Timeline heatmap
                heatmap_fig = self.chart_generator.generate_timeline_heatmap(filtered)
                st.plotly_chart(heatmap_fig, use_container_width=True)

            st.subheader("Event Details")
            events_per_page = 10
            total_pages = (len(filtered) - 1) // events_per_page + 1
            page = st.selectbox("Page", list(range(1, total_pages + 1)))
            start_idx = (page - 1) * events_per_page
            end_idx = min(start_idx + events_per_page, len(filtered))

            for ev in filtered[start_idx:end_idx]:
                title = f"{ev.get('timestamp')} - {ev.get('source', 'unknown')} ({str(ev.get('severity', '')).upper()})"
                with st.expander(title):
                    st.write(f"**Event Type:** {ev.get('event_type', 'N/A')}")
                    st.write(f"**Message:** {ev.get('message', 'N/A')}")
                    if ev.get("event_id"):
                        st.write(f"**Event ID:** {ev.get('event_id')}")
        else:
            st.info("No events match the selected filters or time range.")

    # ---------------------------------------------
    # Reports
    # ---------------------------------------------
    def _render_reports_tab(self):
        st.header("📄 Report Generation")

        if not st.session_state.get("analysis_complete", False):
            st.info("Run an analysis to generate reports")
            return

        report: ForensicReport = st.session_state.get("forensic_report")
        if not report:
            st.error("No report found in session")
            return

        # Report Overview
        st.subheader("📋 Report Overview")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Case ID", str(getattr(report, "case_id", "N/A")))
        with col2:
            st.metric("Total Events", f"{getattr(report, 'total_events', 0):,}")
        with col3:
            sev = (
                str(getattr(getattr(report.analysis.initial_analysis, "severity", ""), "value", "")).upper()
                if getattr(report, "analysis", None) and getattr(report.analysis, "initial_analysis", None)
                else "N/A"
            )
            st.metric("Severity", sev)
        with col4:
            conf = (
                f"{float(getattr(report.analysis.initial_analysis, 'confidence', 0.0)):.1%}"
                if getattr(report, "analysis", None) and getattr(report.analysis, "initial_analysis", None)
                else "N/A"
            )
            st.metric("Confidence", conf)

        # Executive Summary
        st.subheader("📝 Executive Summary")
        st.info(str(getattr(report, "executive_summary", "No summary provided.")))

        # Report Generation Options
        st.markdown("---")
        st.subheader("📤 Download Reports")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("### 📄 Word Document Report")
            st.markdown("""
            **Includes:**
            - Executive summary
            - Key metrics and findings
            - IOCs and MITRE mappings
            - Timeline analysis
            - Recommendations
            - Professional formatting
            """)
            
            if Document is not None:
                try:
                    docx_bytes = self.report_generator.generate_docx(report)
                    st.download_button(
                        label="📥 Download Word Report (.docx)",
                        data=docx_bytes,
                        file_name=f"forensic_report_{getattr(report, 'case_id', 'case')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.error(f"Failed to generate DOCX: {e}")
            else:
                st.warning("python-docx not installed. Cannot generate Word report.")
        
        with col2:
            st.markdown("### 📊 JSON Data Export")
            st.markdown("""
            **Includes:**
            - Complete analysis data
            - Raw findings and metadata
            - Machine-readable format
            - For integration with other tools
            """)
            
            json_bytes = self.report_generator.generate_json(report)
            st.download_button(
                label="📥 Download JSON Data",
                data=json_bytes,
                file_name=f"forensic_data_{getattr(report, 'case_id', 'case')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True,
            )
        
        with col3:
            st.markdown("### 🌐 HTML Report")
            st.markdown("""
            **Includes:**
            - Professional web-based report
            - Interactive elements
            - Responsive design
            - Easy to share and view
            """)
            
            html_content = self.report_generator.generate_html_report(report)
            st.download_button(
                label="📥 Download HTML Report",
                data=html_content.encode('utf-8'),
                file_name=f"forensic_report_{getattr(report, 'case_id', 'case')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                mime="text/html",
                use_container_width=True,
            )

        # Report Preview
        st.markdown("---")
        st.subheader("👀 Report Preview")
        
        # Tabs for different preview sections
        preview_tab1, preview_tab2, preview_tab3 = st.tabs(["Key Findings", "IOCs & MITRE", "Raw Data"])
        
        with preview_tab1:
            st.markdown("### Key Findings")
            
            # IOCs Summary
            iocs = getattr(getattr(report.analysis, 'initial_analysis', {}), 'iocs', []) if getattr(report, 'analysis', None) else []
            if iocs:
                st.markdown("**Indicators of Compromise:**")
                ioc_data = []
                for ioc in iocs:
                    ioc_data.append({
                        "Type": str(getattr(ioc, "type", "")).upper(),
                        "Value": str(getattr(ioc, "value", "")),
                        "Confidence": f"{float(getattr(ioc, 'confidence', 0.0)):.1%}" if getattr(ioc, "confidence", None) is not None else "N/A",
                        "Context": str(getattr(ioc, "context", "")) or "N/A",
                    })
                df_iocs = pd.DataFrame(ioc_data)
                st.dataframe(df_iocs, use_container_width=True)
            else:
                st.info("No IOCs identified")
            
            # Recommendations
            recs = getattr(getattr(report, "analysis", {}), "recommendations", [])
            if recs:
                st.markdown("**Recommendations:**")
                for i, r in enumerate(recs, 1):
                    st.write(f"{i}. {r}")
            else:
                st.info("No specific recommendations generated")
        
        with preview_tab2:
            st.markdown("### MITRE ATT&CK Mappings")
            mitre = getattr(getattr(report.analysis, 'initial_analysis', {}), 'mitre_mappings', []) if getattr(report, 'analysis', None) else []
            if mitre:
                mitre_rows = []
                for m in mitre:
                    mitre_rows.append({
                        "Tactic ID": getattr(m, "tactic_id", ""),
                        "Tactic": getattr(m, "tactic_name", ""),
                        "Technique ID": getattr(m, "technique_id", "N/A"),
                        "Technique": getattr(m, "technique_name", "N/A"),
                        "Confidence": f"{float(getattr(m, 'confidence', 0.0)):.1%}" if getattr(m, "confidence", None) is not None else "N/A",
                    })
                df_mitre = pd.DataFrame(mitre_rows)
                st.dataframe(df_mitre, use_container_width=True)
            else:
                st.info("No MITRE ATT&CK mappings identified")
        
        with preview_tab3:
            st.markdown("### Raw Analysis Data")
            with st.expander("Complete JSON Data", expanded=False):
                st.code(_pretty_json(report), language="json")

        # Chart Export (if enabled)
        if st.session_state.get("enable_charts"):
            st.markdown("---")
            st.subheader("📈 Chart Export")
            st.info("Chart export functionality coming soon! This will allow you to download all visualizations as images or interactive HTML files.")

    # ---------------------------------------------
    # Semantic Search
    # ---------------------------------------------
    def _render_semantic_search_tab(self):
        st.header("Semantic Search")
        st.markdown("---")

        # Search source information
        st.subheader("Search Sources")
        col1, col2 = st.columns(2)
        
        with col1:
            st.info("**ChromaDB**: Your uploaded security logs with AI-generated embeddings")
            if st.session_state.get("stored_logs_count"):
                st.success(f"**{st.session_state['stored_logs_count']} logs** available for search")
            else:
                st.warning("**No logs stored yet**. Upload and analyze logs first to enable semantic search.")
        
        with col2:
            st.info("**Local Knowledge Base**: Pre-populated cybersecurity knowledge (MITRE, attack patterns)")
            st.success("**5 knowledge documents** available")

        st.markdown("---")

        # Manual log storage for testing
        st.subheader("Store Logs Manually (Testing)")
        
        # Test ChromaDB connection
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Test ChromaDB Connection"):
                try:
                    from chromadb_logs import search_logs
                    # Try a simple search to test connection
                    results = search_logs("test", top_k=1)
                    st.success("ChromaDB connection working!")
                    st.info(f"Found {len(results)} existing logs")
                except Exception as e:
                    st.error(f"ChromaDB connection failed: {str(e)[:100]}...")
                    logger.exception("ChromaDB connection test failed")
        
        with col2:
            if st.button("Store Sample Security Logs"):
                try:
                    from chromadb_logs import ingest_logs
                    
                    sample_logs = [
                        "2025-01-01 10:00:00 INFO User login succeeded for alice from 10.0.0.1",
                        "2025-01-01 10:05:10 WARN Multiple failed SSH attempts detected from 203.0.113.5",
                        "2025-01-01 11:22:45 ERROR PowerShell execution blocked by AppLocker on host WIN10-01",
                        "2025-01-01 12:15:30 WARN Unusual network activity detected from internal host 192.168.1.100",
                        "2025-01-01 13:45:12 INFO Database backup completed successfully",
                        "2025-01-01 14:20:00 ERROR Authentication failure for admin account from 198.51.100.23",
                        "2025-01-01 15:30:45 WARN Large data export initiated by user john.doe",
                        "2025-01-01 16:12:18 INFO Firewall rule updated to block suspicious IP 185.220.101.182"
                    ]
                    
                    with st.spinner("Storing sample logs with embeddings..."):
                        stored_items = ingest_logs(sample_logs)
                        st.session_state["stored_logs_count"] = len(stored_items)
                        st.success(f"Stored {len(stored_items)} sample logs in ChromaDB")
                        
                except Exception as e:
                    st.error(f"Failed to store sample logs: {e}")
                    logger.exception("Sample log storage error")

        # Real semantic search interface
        st.subheader("Search Your Logs")
        
        if not st.session_state.get("stored_logs_count", 0):
            st.warning("No logs available for search. Please upload and analyze logs first or use the manual storage above.")
            return

        # Search input
        search_query = st.text_input(
            "Enter search query", 
            placeholder="e.g., 'failed login', 'suspicious PowerShell', 'network anomaly'",
            help="Search for semantically similar security events in your uploaded logs"
        )

        if search_query:
            try:
                from chromadb_logs import search_logs
                
                with st.spinner("Searching logs with AI embeddings..."):
                    # Search ChromaDB for similar logs
                    results = search_logs(search_query, top_k=10)
                
                if results:
                    st.success(f"Found {len(results)} relevant log entries")
                    
                    # Display results with similarity scores
                    for i, result in enumerate(results, 1):
                        with st.expander(f"Result {i} - Score: {result.get('score', 0):.3f}"):
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                st.write(f"**Log ID:** {result.get('logId', 'N/A')}")
                                st.write(f"**Text:** {result.get('logText', 'N/A')}")
                                st.write(f"**Confidence:** {result.get('confidence', 'N/A')}")
                            with col2:
                                score = result.get('score', 0)
                                if score > 0.8:
                                    st.success("High Match")
                                elif score > 0.6:
                                    st.warning("Medium Match")
                                else:
                                    st.info("Low Match")
                    
                    # Summary statistics
                    st.subheader("Search Summary")
                    high_matches = len([r for r in results if r.get('score', 0) > 0.8])
                    medium_matches = len([r for r in results if 0.6 < r.get('score', 0) <= 0.8])
                    low_matches = len([r for r in results if r.get('score', 0) <= 0.6])
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("High Matches", high_matches)
                    with col2:
                        st.metric("Medium Matches", medium_matches)
                    with col3:
                        st.metric("Low Matches", low_matches)
                        
                else:
                    st.info("No relevant logs found for your query. Try different keywords or check if logs contain similar content.")
                    
            except Exception as e:
                st.error(f"Search failed: {e}")
                logger.exception("Semantic search error")

        # Local knowledge base search
        st.markdown("---")
        st.subheader("Search Cybersecurity Knowledge Base")
        
        kb_query = st.text_input(
            "Knowledge base query",
            placeholder="e.g., 'MITRE T1078', 'lateral movement', 'PowerShell attacks'",
            help="Search pre-populated cybersecurity knowledge and attack patterns"
        )
        
        if kb_query:
            try:
                from rag_enrichment import LocalKnowledgeBase
                
                kb = LocalKnowledgeBase()
                kb_results = kb.query(kb_query, n_results=3)
                
                if kb_results:
                    st.success(f"Found {len(kb_results)} knowledge base entries")
                    
                    for i, result in enumerate(kb_results, 1):
                        with st.expander(f"Knowledge {i} - Relevance: {1 - result.get('distance', 0):.3f}"):
                            st.write(f"**Content:** {result.get('content', 'N/A')}")
                            st.write(f"**Type:** {result.get('metadata', {}).get('type', 'N/A')}")
                            if result.get('metadata', {}).get('technique_id'):
                                st.write(f"**Technique ID:** {result['metadata']['technique_id']}")
                else:
                    st.info("No relevant knowledge base entries found.")
                    
            except Exception as e:
                st.error(f"Knowledge base search failed: {e}")
                logger.exception("Knowledge base search error")

    # ---------------------------------------------
    # Attack Chain Rendering
    # ---------------------------------------------
    def _render_attack_chain(self, chain: Any):
        """Render attack chain using PyVis if available; otherwise fallback to table.

        Expected structures supported:
          • List[dict]: items with keys like {from, to, label, severity, weight, timestamp}
          • List[str]: sequential step names
        """
        # Normalize items
        if isinstance(chain, list) and chain and isinstance(chain[0], dict):
            edges = chain
            nodes = set()
            for e in edges:
                nodes.add(e.get("from", ""))
                nodes.add(e.get("to", ""))
        elif isinstance(chain, list) and chain and isinstance(chain[0], str):
            edges = [
                {"from": chain[i], "to": chain[i + 1], "label": "next"}
                for i in range(len(chain) - 1)
            ]
            nodes = set(chain)
        else:
            st.dataframe(pd.DataFrame(_safe_to_dict(chain)))
            return

        # Try interactive graph with PyVis
        if Network is not None:
            try:
                net = Network(height="520px", width="100%", bgcolor="#FFFFFF", font_color="#222222")
                net.barnes_hut()

                # Add nodes
                for n in nodes:
                    if not n:
                        continue
                    net.add_node(str(n), label=str(n))

                # Add edges
                for e in edges:
                    src = str(e.get("from", ""))
                    dst = str(e.get("to", ""))
                    lbl = str(e.get("label", ""))
                    if src and dst:
                        net.add_edge(src, dst, title=lbl, label=lbl)

                with tempfile.NamedTemporaryFile(delete=False, suffix="_attack_chain.html") as tmp:
                    net.write_html(tmp.name, open_browser=False, notebook=False)
                    html = Path(tmp.name).read_text(encoding="utf-8")
                components.html(html, height=540, scrolling=True)
                return
            except Exception:  # pragma: no cover
                logger.exception("PyVis render failed, falling back to table")

        # Fallback
        st.dataframe(pd.DataFrame(edges))
    
    def _render_saved_analysis(self, analysis_id: int):
        """Render a saved analysis from the database."""
        try:
            from results_storage import get_results_storage
            storage = get_results_storage()
            
            analysis = storage.get_analysis_by_id(analysis_id)
            if not analysis:
                st.error(f"Analysis with ID {analysis_id} not found")
                return
            
            # Header with back button
            col1, col2 = st.columns([3, 1])
            with col1:
                st.header(f"Saved Analysis: {analysis['case_name']}")
            with col2:
                if st.button("← Back to Saved Results"):
                    st.session_state.pop("view_analysis_id", None)
                    st.rerun()
            
            st.markdown("---")
            
            # Analysis metadata
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Case ID", analysis['case_id'])
            with col2:
                st.metric("Analysis Date", analysis['created_at'][:19].replace("T", " "))
            with col3:
                st.metric("Files Analyzed", analysis['file_count'])
            with col4:
                st.metric("Log Entries", f"{analysis['log_entries_count']:,}")
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                severity = analysis['severity'].upper() if analysis['severity'] else "UNKNOWN"
                st.metric("Severity", severity)
            with col2:
                confidence = f"{analysis['confidence']:.1%}" if analysis['confidence'] else "N/A"
                st.metric("Confidence", confidence)
            with col3:
                st.metric("IOCs Found", analysis['iocs_count'])
            with col4:
                st.metric("MITRE Mappings", analysis['mitre_mappings_count'])
            
            # Executive Summary
            st.subheader("📋 Executive Summary")
            st.info(analysis['executive_summary'])
            
            # Detailed Results
            detailed_results = analysis.get('detailed_results', {})
            
            # IOCs
            if detailed_results.get('iocs'):
                st.subheader("Indicators of Compromise")
                ioc_data = []
                for ioc in detailed_results['iocs']:
                    ioc_data.append({
                        "Type": str(ioc.get('type', '')).upper(),
                        "Value": str(ioc.get('value', '')),
                        "Confidence": f"{ioc.get('confidence', 0):.1%}" if ioc.get('confidence') else "N/A",
                        "Context": str(ioc.get('context', '')) or "N/A"
                    })
                
                if ioc_data:
                    df_iocs = pd.DataFrame(ioc_data)
                    st.dataframe(df_iocs, width='stretch')
                else:
                    st.info("No IOCs identified")
            
            # MITRE ATT&CK
            if detailed_results.get('mitre'):
                st.subheader("MITRE ATT&CK Mapping")
                mitre_data = []
                for mapping in detailed_results['mitre']:
                    mitre_data.append({
                        "Tactic ID": mapping.get('tactic_id', ''),
                        "Tactic": mapping.get('tactic_name', ''),
                        "Technique ID": mapping.get('technique_id', 'N/A'),
                        "Technique": mapping.get('technique_name', 'N/A'),
                        "Confidence": f"{mapping.get('confidence', 0):.1%}" if mapping.get('confidence') else "N/A"
                    })
                
                if mitre_data:
                    df_mitre = pd.DataFrame(mitre_data)
                    st.dataframe(df_mitre, width='stretch')
                else:
                    st.info("No MITRE ATT&CK mappings identified")
            
            # Recommendations
            if detailed_results.get('recommendations'):
                st.subheader("💡 Recommendations")
                for i, rec in enumerate(detailed_results['recommendations'], 1):
                    st.write(f"{i}. {rec}")
            else:
                st.subheader("💡 Recommendations")
                st.info("No specific recommendations generated")
            
            # Timeline
            if detailed_results.get('timeline'):
                st.subheader("⏰ Event Timeline")
                timeline_data = []
                for event in detailed_results['timeline'][:100]:  # Show first 100 events
                    timeline_data.append({
                        "Timestamp": event.get('timestamp', ''),
                        "Source": event.get('source', ''),
                        "Severity": event.get('severity', ''),
                        "Event Type": event.get('event_type', ''),
                        "Message": event.get('message', '')[:100] + "..." if len(str(event.get('message', ''))) > 100 else event.get('message', '')
                    })
                
                if timeline_data:
                    df_timeline = pd.DataFrame(timeline_data)
                    st.dataframe(df_timeline, width='stretch')
                    
                    if len(detailed_results['timeline']) > 100:
                        st.info(f"Showing first 100 of {len(detailed_results['timeline'])} timeline events")
                else:
                    st.info("No timeline data available")
            
            # File Information
            if analysis.get('files'):
                st.subheader("Analyzed Files")
                file_data = []
                for file_info in analysis['files']:
                    file_data.append({
                        "Name": file_info['original_name'],
                        "Type": file_info['file_type'],
                        "Category": file_info['file_category'],
                        "Size": f"{file_info['file_size'] / (1024*1024):.2f} MB"
                    })
                
                df_files = pd.DataFrame(file_data)
                st.dataframe(df_files, width='stretch')
            
            # Analysis Configuration
            if analysis.get('analysis_config'):
                st.subheader("Analysis Configuration")
                config = analysis['analysis_config']
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Processing Mode:** {config.get('processing_mode', 'N/A')}")
                    st.write(f"**Skip ChromaDB:** {config.get('skip_chromadb', 'N/A')}")
                with col2:
                    st.write(f"**Max Entries:** {config.get('max_entries', 'N/A')}")
                    st.write(f"**Analysis Depth:** {config.get('analysis_depth', 'N/A')}")
            
            # Action buttons
            st.markdown("---")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("📥 Download Report", key=f"download_saved_{analysis_id}"):
                    st.info("Download functionality coming soon!")
            
            with col2:
                if st.button("🔄 Re-analyze", key=f"reanalyze_saved_{analysis_id}"):
                    st.info("Re-analysis functionality coming soon!")
            
            with col3:
                if st.button("Delete Analysis", key=f"delete_saved_{analysis_id}"):
                    if storage.delete_analysis(analysis_id):
                        st.success("Analysis deleted successfully!")
                        st.session_state.pop("view_analysis_id", None)
                        st.rerun()
                    else:
                        st.error("Failed to delete analysis")
                        
        except Exception as e:
            st.error(f"Error loading saved analysis: {e}")
            logger.exception("Error in saved analysis view")

    def _update_file_history(self, uploaded_files, case_id: str, log_count: int):
        """Update file history with newly uploaded files."""
        if "file_history" not in st.session_state:
            st.session_state["file_history"] = []
        
        current_time = datetime.now()
        
        for uploaded_file in uploaded_files:
            file_info = {
                "id": f"file_{current_time.timestamp()}_{uploaded_file.name}",
                "filename": uploaded_file.name,
                "size_bytes": len(uploaded_file.getbuffer()),
                "size_mb": len(uploaded_file.getbuffer()) / (1024 * 1024),
                "file_type": uploaded_file.type or "Unknown",
                "upload_time": current_time.isoformat(),
                "case_id": case_id or "No Case ID",
                "log_entries_processed": log_count,
                "analysis_completed": True,
                "status": "Completed"
            }
            
            # Check if file already exists in history (by name and size)
            existing_file = next(
                (f for f in st.session_state["file_history"] 
                 if f["filename"] == file_info["filename"] and 
                 abs(f["size_bytes"] - file_info["size_bytes"]) < 100),  # Allow 100 byte difference
                None
            )
            
            if existing_file:
                # Update existing entry
                existing_file.update({
                    "upload_time": current_time.isoformat(),
                    "case_id": case_id or "No Case ID",
                    "log_entries_processed": log_count,
                    "analysis_completed": True,
                    "status": "Re-analyzed"
                })
            else:
                # Add new entry
                st.session_state["file_history"].append(file_info)
        
        # Keep only last 50 files to prevent memory issues
        if len(st.session_state["file_history"]) > 50:
            st.session_state["file_history"] = st.session_state["file_history"][-50:]
        
        logger.info(f"Updated file history with {len(uploaded_files)} files")

    def _get_file_history(self) -> List[Dict[str, Any]]:
        """Get the current file history."""
        return st.session_state.get("file_history", [])

    def _render_saved_results_tab(self):
        """Render the saved results tab."""
        st.header("Saved Analysis Results")
        st.markdown("---")
        
        try:
            from results_storage import get_results_storage
            storage = get_results_storage()
            
            # Get statistics
            stats = storage.get_statistics()
            
            # Display statistics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Analyses", stats.get('total_analyses', 0))
            with col2:
                st.metric("Recent (30 days)", stats.get('recent_analyses', 0))
            with col3:
                st.metric("Total Files", stats.get('total_files', 0))
            with col4:
                st.metric("Total IOCs", stats.get('total_iocs', 0))
            
            # Search and filter controls
            st.markdown("---")
            col1, col2, col3 = st.columns([2, 1, 1])
            
            with col1:
                search_query = st.text_input(
                    "Search analyses",
                    placeholder="Search by case ID, case name, or content...",
                    help="Search through saved analysis results"
                )
            
            with col2:
                if st.button("Search"):
                    st.session_state["search_results"] = storage.search_analyses(search_query) if search_query else []
                    st.session_state["show_search_results"] = True
            
            with col3:
                if st.button("🔄 Refresh"):
                    st.rerun()
            
            # Display results
            if st.session_state.get("show_search_results") and st.session_state.get("search_results"):
                st.subheader(f"Search Results ({len(st.session_state['search_results'])} found)")
                analyses = st.session_state["search_results"]
            else:
                st.subheader("📋 Recent Analyses")
                analyses = storage.get_all_analyses(limit=20)
            
            if analyses:
                # Create DataFrame for display
                display_data = []
                for analysis in analyses:
                    display_data.append({
                        "ID": analysis['id'],
                        "Case ID": analysis['case_id'],
                        "Case Name": analysis['case_name'],
                        "Date": analysis['created_at'][:19].replace("T", " "),
                        "Files": analysis['file_count'],
                        "Log Entries": f"{analysis['log_entries_count']:,}",
                        "Severity": analysis['severity'].upper() if analysis['severity'] else "UNKNOWN",
                        "Confidence": f"{analysis['confidence']:.1%}" if analysis['confidence'] else "N/A",
                        "IOCs": analysis['iocs_count'],
                        "MITRE": analysis['mitre_mappings_count']
                    })
                
                df_analyses = pd.DataFrame(display_data)
                st.dataframe(df_analyses, width='stretch', hide_index=True)
                
                # Analysis details
                st.subheader("📄 Analysis Details")
                
                for analysis in analyses:
                    with st.expander(f"{analysis['case_name']} (ID: {analysis['id']}) - {analysis['created_at'][:19].replace('T', ' ')}"):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.write(f"**Case ID:** {analysis['case_id']}")
                            st.write(f"**Analysis Date:** {analysis['created_at'][:19].replace('T', ' ')}")
                            st.write(f"**Files Analyzed:** {analysis['file_count']}")
                            st.write(f"**Log Entries:** {analysis['log_entries_count']:,}")
                            st.write(f"**Severity:** {analysis['severity'].upper() if analysis['severity'] else 'UNKNOWN'}")
                            st.write(f"**Confidence:** {analysis['confidence']:.1%}" if analysis['confidence'] else "**Confidence:** N/A")
                        
                        with col2:
                            st.write(f"**IOCs Found:** {analysis['iocs_count']}")
                            st.write(f"**MITRE Mappings:** {analysis['mitre_mappings_count']}")
                            st.write(f"**Recommendations:** {analysis['recommendations_count']}")
                            if analysis.get('zip_file_path'):
                                st.write(f"**ZIP File:** {analysis['zip_file_path']}")
                            st.write(f"**Status:** {analysis['status']}")
                        
                        # Executive summary
                        st.markdown("---")
                        st.write("**Executive Summary:**")
                        st.info(analysis['executive_summary'])
                        
                        # Action buttons
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            if st.button(f"View Full Report", key=f"view_{analysis['id']}"):
                                st.session_state["view_analysis_id"] = analysis['id']
                                st.rerun()
                        
                        with col2:
                            if st.button(f"📥 Download Report", key=f"download_{analysis['id']}"):
                                # TODO: Implement download functionality
                                st.info("Download functionality coming soon!")
                        
                        with col3:
                            if st.button(f"Delete", key=f"delete_{analysis['id']}"):
                                if storage.delete_analysis(analysis['id']):
                                    st.success(f"Analysis {analysis['id']} deleted successfully!")
                                    st.rerun()
                                else:
                                    st.error("Failed to delete analysis")
            
            else:
                st.info("📭 No saved analyses found. Run an analysis first to see results here.")
                
                # Show example of what will be saved
                st.markdown("---")
                st.subheader("💡 What gets saved?")
                st.markdown("""
                When you run an analysis, the following information is automatically saved:
                - **Case details** (ID, name, date)
                - **Analysis results** (IOCs, MITRE mappings, recommendations)
                - **File information** (names, types, categories)
                - **Timeline data** (event sequences)
                - **Configuration** (analysis settings used)
                """)
                
        except Exception as e:
            st.error(f"Error loading saved results: {e}")
            logger.exception("Error in saved results tab")
    
    def _render_file_history_tab(self):
        """Render the file history tab."""
        st.header("File History")
        st.markdown("---")
        
        # File history controls
        col1, col2, col3 = st.columns([2, 1, 1])
        
        with col1:
            st.subheader("Previously Uploaded Files")
        
        with col2:
            if st.button("Refresh History"):
                st.rerun()
        
        with col3:
            if st.button("Clear History"):
                if st.session_state.get("file_history"):
                    st.session_state["file_history"] = []
                    st.success("File history cleared!")
                    st.rerun()
        
        # Get file history
        file_history = self._get_file_history()
        
        if not file_history:
            st.info("No files have been uploaded yet. Upload files in the 'Upload & Analyze' tab to see them here.")
            return
        
        # File history statistics
        st.subheader("Statistics")
        total_files = len(file_history)
        total_size_mb = sum(f["size_mb"] for f in file_history)
        unique_cases = len(set(f["case_id"] for f in file_history))
        total_logs = sum(f["log_entries_processed"] for f in file_history)
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Files", total_files)
        with col2:
            st.metric("Total Size", f"{total_size_mb:.1f} MB")
        with col3:
            st.metric("Unique Cases", unique_cases)
        with col4:
            st.metric("Total Log Entries", f"{total_logs:,}")
        
        # File filtering
        st.markdown("---")
        st.subheader("Filter Files")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Filter by case ID
            case_ids = ["All Cases"] + sorted(list(set(f["case_id"] for f in file_history)))
            selected_case = st.selectbox("Case ID Filter", case_ids)
        
        with col2:
            # Filter by file type
            file_types = ["All Types"] + sorted(list(set(f["file_type"] for f in file_history)))
            selected_type = st.selectbox("File Type Filter", file_types)
        
        with col3:
            # Filter by status
            statuses = ["All Statuses"] + sorted(list(set(f["status"] for f in file_history)))
            selected_status = st.selectbox("Status Filter", statuses)
        
        # Apply filters
        filtered_files = file_history.copy()
        
        if selected_case != "All Cases":
            filtered_files = [f for f in filtered_files if f["case_id"] == selected_case]
        
        if selected_type != "All Types":
            filtered_files = [f for f in filtered_files if f["file_type"] == selected_type]
        
        if selected_status != "All Statuses":
            filtered_files = [f for f in filtered_files if f["status"] == selected_status]
        
        # Display filtered files
        st.markdown("---")
        st.subheader(f"Files ({len(filtered_files)} of {total_files})")
        
        if filtered_files:
            # Sort by upload time (newest first)
            filtered_files.sort(key=lambda x: x["upload_time"], reverse=True)
            
            # Create DataFrame for display
            display_data = []
            for file_info in filtered_files:
                display_data.append({
                    "Filename": file_info["filename"],
                    "Size": f"{file_info['size_mb']:.2f} MB",
                    "Type": file_info["file_type"],
                    "Case ID": file_info["case_id"],
                    "Log Entries": f"{file_info['log_entries_processed']:,}",
                    "Upload Time": file_info["upload_time"][:19].replace("T", " "),
                    "Status": file_info["status"]
                })
            
            df_files = pd.DataFrame(display_data)
            st.dataframe(df_files, width='stretch', hide_index=True)
            
            # File details in expandable sections
            st.subheader("File Details")
            
            for i, file_info in enumerate(filtered_files):
                with st.expander(f"{file_info['filename']} - {file_info['status']}"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write(f"**File ID:** {file_info['id']}")
                        st.write(f"**Size:** {file_info['size_mb']:.2f} MB ({file_info['size_bytes']:,} bytes)")
                        st.write(f"**Type:** {file_info['file_type']}")
                        st.write(f"**Case ID:** {file_info['case_id']}")
                    
                    with col2:
                        st.write(f"**Upload Time:** {file_info['upload_time'][:19].replace('T', ' ')}")
                        st.write(f"**Log Entries:** {file_info['log_entries_processed']:,}")
                        st.write(f"**Analysis Status:** {file_info['status']}")
                        st.write(f"**Analysis Completed:** {'Yes' if file_info['analysis_completed'] else 'No'}")
                    
                    # Action buttons
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if st.button(f"Re-analyze", key=f"reanalyze_{i}"):
                            st.info("Re-analysis feature coming soon! This will allow you to re-analyze previously uploaded files.")
                    
                    with col2:
                        if st.button(f"View Report", key=f"view_report_{i}"):
                            st.info("Report viewing feature coming soon! This will show the analysis report for this file.")
                    
                    with col3:
                        if st.button(f"Remove", key=f"remove_{i}"):
                            # Remove file from history
                            st.session_state["file_history"] = [f for f in st.session_state["file_history"] if f["id"] != file_info["id"]]
                            st.success(f"Removed {file_info['filename']} from history!")
                            st.rerun()
        else:
            st.warning("No files match the selected filters.")
        
        # Export functionality
        st.markdown("---")
        st.subheader("Export History")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if filtered_files:
                # Export filtered files as CSV
                df_export = pd.DataFrame(filtered_files)
                csv_data = df_export.to_csv(index=False)
                st.download_button(
                    label="Download CSV",
                    data=csv_data,
                    file_name=f"file_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    width='stretch'
                )
        
        with col2:
            if filtered_files:
                # Export as JSON
                json_data = json.dumps(filtered_files, indent=2, default=str)
                st.download_button(
                    label="Download JSON",
                    data=json_data,
                    file_name=f"file_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    width='stretch'
                )


# -----------------------------------------------------------------------------
# Entrypoint
# -----------------------------------------------------------------------------
def main() -> None:
    """Run the Streamlit dashboard app."""
    ForensicDashboard().run()


if __name__ == "__main__":
    main()
